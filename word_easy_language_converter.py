import json
import re
import os

import docx
from nltk.tokenize import sent_tokenize
from openai import OpenAI

client = OpenAI(api_key="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")

#models = client.models.list()
#pass

def is_heading(paragraph):
    """Determine if a paragraph is a heading."""
    # List of possible heading styles in German Word
    heading_styles = ['Title', 'Titel', 'Heading', 'Überschrift']
    return any(paragraph.style.name.startswith(heading) for heading in heading_styles)

def split_into_chapters(paragraphs):
    """
    Split the document into chapters based on heading styles.
    """
    chapters = {}
    current_chapter = "Heading 1: Impressum"
    chapters[current_chapter] = []

    for paragraph in paragraphs:
        # Check if the paragraph is a heading
        if is_heading(paragraph):
            if paragraph.text not in chapters:
                chapter_name = "Heading %s: %s" % (paragraph.style.name.split(" ")[-1], paragraph.text)
                chapters[chapter_name] = []
                current_chapter = chapter_name
            # Save the current chapter and start a new one
        else:
            chapters[current_chapter].append(paragraph.text)

    return chapters

def estimate_gpt_token_count(sentence):
    """
    Estimates the GPT token count for a given sentence.
    This approximation assumes 4 tokens per character.

    :param sentence: A sentence whose GPT token count needs to be estimated.
    :return: Estimated token count.
    """
    # Estimate token count: 4 tokens per character
    estimated_token_count = len(sentence) * 4
    return estimated_token_count

def extract_text_from_docx(docx_path):
    """Extract and categorize text from a docx file into headings and corresponding text."""
    doc = docx.Document(docx_path)
    content = {}
    current_heading = "Introduction"  # Default heading
    for para in doc.paragraphs:
        if is_heading(para):
            current_heading = para.text
            content[current_heading] = []
        else:
            content[current_heading].append(para.text)
    return content

def read_word_document(file_path):
    """
    Reads a Word document and extracts the text.

    :param file_path: Path to the Word document file.
    :return: A single string containing all the text from the document.
    """
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_json(text):
    # Find the start of the JSON block
    start_index = text.find("{")
    if start_index == -1:
        return "JSON block not found."

    # Find the end of the JSON block
    end_index = text.rfind("}")
    if end_index == -1:
        return "JSON block not found."

    # Extract and return the JSON block
    return text[start_index:end_index + 1]

def split_text_into_chunks(text, max_tokens, prompt_txt):
    """
    Splits the text into chunks, each having a size close to the specified maximum token count.
    The splitting ensures that sentences are not cut in the middle.

    :param text: The complete text to be split.
    :param max_tokens: Maximum token count for each chunk.
    :return: A list of text chunks.
    """

    pattern = r'%[^%]+%'
    real_prompt_txt = re.sub(pattern, "", prompt_txt)
    prompt_token_count = estimate_gpt_token_count(real_prompt_txt)

    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        # Estimated token count for the sentence
        sentence_token_count = estimate_gpt_token_count(sentence) + prompt_token_count
        # Estimated token count for the current chunk
        current_chunk_token_count = estimate_gpt_token_count(current_chunk)

        # Check if adding this sentence would exceed the max token count
        if current_chunk_token_count + sentence_token_count > max_tokens:
            # If the current chunk is not empty, add it to the chunks list
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ""

        current_chunk += sentence + " "

    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append(current_chunk)

    return chunks

def read_word_document(file_path):
    """
    Read a Word document and return its content as a list of paragraphs.
    """
    doc = docx.Document(file_path)
    return [paragraph for paragraph in doc.paragraphs if len(paragraph.text.replace("\n", "")) > 0]

def get_filename_without_extension(path):
    base_name = os.path.basename(path)
    filename_without_extension, _ = os.path.splitext(base_name)
    return filename_without_extension

# Example usage
docx_path = "//dkg-dc-01/Daten/Daten/Arbeitsverzeichnisse/OL/PatientenLL/Manuskripte PLL/Weichgewebesarkome/2021_11_10_PL_Weichgewebesarkome_final_korr.docx"
prompt_text = '%TEXT%.'
out_filename = get_filename_without_extension(docx_path)+".txt"
content = read_word_document(docx_path)
chapters = split_into_chapters(content)

previous_text = ""

for chapter, paragraphs in chapters.items():
    if 'Impressum' in chapter:
        continue
    text = "\n".join([txt for txt in paragraphs])
    segments = split_text_into_chunks(text, 4096, prompt_text)

    # Now, segments can be used for GPT-4 processing
    for i in range(len(segments)):
        print("%s, Segment %s/%s" % (chapter, i+1, len(segments)))
        segment = segments[i]
        query = 'Transform this text into Easy Language in German: %s\nLIMITIATIONS:1. Clear and simple sentences: Use short and straightforward sentences, Avoid complex sentence structures\n2. Avoidance of technical terms: Replace technical or specialized vocabulary with simple and common words, Provide explanations or synonyms for any necessary technical terms.\n3. Limited word count: Keep texts concise and avoid unnecessary details. Use only essential information to convey the message.\n4. Use of everyday words: Utilize common, everyday vocabulary that is familiar to the target audience, Avoid idiomatic expressions or colloquialisms.\n5. Precise and clear wording: Ensure that the content is accurate and unambiguous, Avoid vague or ambiguous language\n6. Active Voice: Prefer using the active voice over passive voice constructions, Make it clear who is performing the action.\n7. Structure and headings: Organize content with clear headings and subheadings, Use bullet points or numbered lists for enumerations\n8. Repetition of key information: Repeat important information throughout the text to reinforce comprehension, Use summaries or key points if necessary. Do not repeat the limitations in your response.' % segment
        error = ''
        success = False

        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": query,
                },
            ],
            model="gpt-4")
        response = chat_completion.choices[0].message.content
        f = open(out_filename, "a", encoding="utf-8")
        f.write("\n%s\n%s" % (chapter, response))
        f.close()