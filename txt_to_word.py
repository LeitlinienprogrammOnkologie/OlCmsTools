import docx

txt_source_path = "2021_11_10_PL_Weichgewebesarkome_final_korr.txt"

doc = docx.Document()

heading_list = []

lines = None
with open(txt_source_path, "r", encoding="utf-8") as f:
    lines = f.readlines()

for line in lines:
    if len(line.strip().strip("\r\n")) == 0:
        continue

    if "Heading" in line:
        if line in heading_list:
            continue
        heading_arr = line.split(":")
        heading_lvl = heading_arr[0].split(" ")[1]
        heading = doc.add_paragraph(heading_arr[1].strip())
        heading.style = "Heading %s" % heading_lvl
        heading_list.append(line)
    else:
        paragraph_arr = line.split("\n")
        for paragraph in paragraph_arr:
            paragraph = paragraph.strip().strip("\n\r")
            if len(paragraph) == 0:
                continue
            if paragraph[0] == "-" or paragraph[0] == "*":
                p = doc.add_paragraph(paragraph.replace("%s " % paragraph[0], ""))
                p.style = "List Bullet"
            else:
                p = doc.add_paragraph(paragraph)


doc.save("test_pll_easy.docx")